import io
from .base import BaseParser

class DOCXParser(BaseParser):
    def parse(self, file_content: bytes) -> dict:
        text = self.extract_text(file_content)
        meta = self.extract_metadata(file_content)
        tables = self.extract_tables(file_content)
        images = self.extract_images(file_content)
        structure = self.extract_structure(file_content)

        words = len(text.split())
        chars = len(text)
        
        return {
            "title": meta.get("title", ""),
            "metadata": meta,
            "sections": structure,
            "paragraphs": [{"text": p.strip()} for p in text.split("\n\n") if p.strip()],
            "tables": tables,
            "images": images,
            "links": [],
            "language": "en",
            "statistics": {
                "word_count": words,
                "character_count": chars,
                "page_count": 1,
                "table_count": len(tables),
                "image_count": len(images)
            }
        }

    def extract_text(self, file_content: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text])
        except ImportError:
            return "python-docx not installed. Return empty fallback."
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"

    def extract_tables(self, file_content: bytes) -> list[dict]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            tables_list = []
            for t_idx, table in enumerate(doc.tables):
                grid = []
                for row in table.rows:
                    grid.append([cell.text.strip() for cell in row.cells])
                tables_list.append({
                    "table_index": t_idx + 1,
                    "data": grid
                })
            return tables_list
        except Exception:
            return []

    def extract_images(self, file_content: bytes) -> list[dict]:
        # python-docx stores images in doc.inline_shapes
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            images_list = []
            for shape_idx, shape in enumerate(doc.inline_shapes):
                images_list.append({
                    "image_index": shape_idx + 1,
                    "width": shape.width,
                    "height": shape.height
                })
            return images_list
        except Exception:
            return []

    def extract_metadata(self, file_content: bytes) -> dict:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            core_properties = doc.core_properties
            return {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "created": str(core_properties.created) if core_properties.created else ""
            }
        except Exception:
            return {}

    def extract_structure(self, file_content: bytes) -> list[dict]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            structure = []
            for p in doc.paragraphs:
                if p.style.name.startswith("Heading"):
                    structure.append({
                        "level": p.style.name.replace("Heading ", ""),
                        "title": p.text
                    })
            return structure
        except Exception:
            return []
